"""Load brand tone examples for copy generation prompts."""

from __future__ import annotations

from pathlib import Path

from app.config import get_settings

BRAND_FILE_EXTENSIONS = {".txt", ".md", ".html", ".docx"}
STYLE_GUIDE_NAMES = ("style-guide.txt", "style-guide.md", "style-guide.docx")
WORD_TEMP_PREFIX = "~$"


def _is_brand_example_file(path: Path) -> bool:
    if not path.is_file():
        return False
    if path.name.startswith(WORD_TEMP_PREFIX):
        return False
    if path.suffix.lower() not in BRAND_FILE_EXTENSIONS:
        return False
    if path.name in STYLE_GUIDE_NAMES:
        return False
    return True


def read_brand_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".html"}:
        return path.read_text(encoding="utf-8").strip()
    if suffix == ".docx":
        return _read_docx(path)
    return ""


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from exc

    doc = Document(path)
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts).strip()


def _load_style_guide(examples_dir: Path) -> tuple[str, list[str]]:
    notes: list[str] = []
    for name in STYLE_GUIDE_NAMES:
        path = examples_dir / name
        if not path.exists():
            continue
        try:
            text = read_brand_file(path)
            if text:
                notes.append(f"Loaded {name}")
                return text, notes
        except Exception as exc:
            notes.append(f"Could not read {name}: {exc}")
    return "", notes


def load_brand_context(max_examples: int | None = None) -> tuple[str, list[str]]:
    settings = get_settings()
    examples_dir = settings.brand_examples_dir
    limit = max_examples or settings.max_brand_examples
    notes: list[str] = []

    style_guide, guide_notes = _load_style_guide(examples_dir)
    notes.extend(guide_notes)

    examples: list[str] = []
    if examples_dir.exists():
        for path in sorted(examples_dir.iterdir()):
            if not _is_brand_example_file(path):
                continue
            try:
                text = read_brand_file(path)
            except Exception as exc:
                notes.append(f"Skipped {path.name}: {exc}")
                continue
            if len(text) > 100:
                examples.append(text[:3000])
                notes.append(f"Loaded {path.name}")
            if len(examples) >= limit:
                break

    if examples:
        notes.append(f"Using {len(examples)} brand example(s) in prompt")

    parts = []
    if style_guide:
        parts.append(f"BRAND STYLE RULES:\n{style_guide}")
    if examples:
        parts.append("EXAMPLE PRODUCT COPY FROM OUR STORE (match this tone):\n")
        for i, ex in enumerate(examples, 1):
            parts.append(f"--- Example {i} ---\n{ex}\n")

    return "\n".join(parts), notes
